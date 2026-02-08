from pdf2image import convert_from_path
import pytesseract
from docx import Document
from PIL import Image
import re
import os

# Set Tesseract path (Windows only)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

PDF_PATH = "result.pdf"
OUTPUT_DOC = "Student_Results.docx"


def crop_student_blocks(image):
    """
    Detect ERN lines and crop student result areas.
    """
    ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    crops = []

    for i, text in enumerate(ocr_data["text"]):
        # Look for ERN pattern like (MU0341....)
        if re.search(r'\(MU.*?\)', text):
            x = ocr_data["left"][i]
            y = ocr_data["top"][i]

            # Crop region around ERN
            # Adjust these values if needed
            crop_box = (
                0,
                max(0, y - 80),         # top
                image.width,
                min(image.height, y + 220)  # bottom
            )

            crop = image.crop(crop_box)
            crops.append(crop)

    return crops


def main():
    images = convert_from_path(PDF_PATH, dpi=300)

    doc = Document()
    doc.add_heading("Student Results", level=1)

    count = 0

    for page_img in images:
        crops = crop_student_blocks(page_img)

        for crop in crops:
            temp_path = f"temp_{count}.png"
            crop.save(temp_path)

            doc.add_picture(temp_path, width=None)
            doc.add_paragraph("\n")

            os.remove(temp_path)
            count += 1

    doc.save(OUTPUT_DOC)
    print(f"Saved {count} student results to {OUTPUT_DOC}")


if __name__ == "__main__":
    main()